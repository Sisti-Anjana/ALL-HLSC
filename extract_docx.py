import os
from docx import Document

docx_files = [
    "DETAILED-PROJECT-STRUCTURE-AND-ARCHITECTURE.docx",
    "PROJECT-TIMELINE-AND-PHASES.docx",
    "UI-DESIGN-AND-MOCKUPS.docx",
    "UI-DESIGN-AND-MOCKUPSs.docx"
]

for docx_file in docx_files:
    if os.path.exists(docx_file):
        try:
            doc = Document(docx_file)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Save to text file
            output_file = docx_file.replace('.docx', '.txt')
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            print(f"Extracted: {docx_file} -> {output_file}")
        except Exception as e:
            print(f"Error processing {docx_file}: {e}")
    else:
        print(f"File not found: {docx_file}")

print("\nExtraction complete!")




















