"""
Convert Markdown to Word Document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            p = doc.add_heading(line[7:], level=6)
        # Code blocks
        elif line.startswith('```'):
            continue  # Skip code block markers
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
        # Bold text
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        # Regular paragraph
        else:
            p = doc.add_paragraph(line)
    
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    markdown_to_docx('COMPLETE_APPLICATION_DOCUMENTATION.md', 'COMPLETE_APPLICATION_DOCUMENTATION.docx')

